"""
Creates TerraAlert_Setup_Guide.docx — a clear step-by-step guide for
running the TerraAlert project on any Windows or macOS / Linux PC.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────────────
def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def note(text, color="1F4E79"):
    p = doc.add_paragraph()
    run = p.add_run("ℹ  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.italic = True
    return p

def warn(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x25)
    run.bold = True
    return p

def code_block(lines: list[str]):
    """Each string in lines becomes a monospace paragraph."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        # light grey background
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        rPr.append(shd)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(11)
    return p

def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.runs[0].font.size = Pt(11)
    return p

# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("TerraAlert", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Project Setup & Run Guide")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub2 = doc.add_paragraph("Early Disaster Warning System — Developer Handover Document")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].font.size = Pt(12)
sub2.runs[0].italic = True

doc.add_paragraph()
note("Follow every step carefully. The project will be running locally in your browser within roughly 5 minutes.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
h("1. System Requirements")
# ═════════════════════════════════════════════════════════════════════════════
body("Before you begin, make sure the following software is installed on your PC:")
bullet("Python 3.10 or later  →  https://www.python.org/downloads/")
bullet("Git (optional, only needed if cloning from a repository)  →  https://git-scm.com")
bullet("A modern browser: Chrome, Firefox, or Edge")
note("During Python installation on Windows, tick the checkbox 'Add Python to PATH' before clicking Install.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("2. Getting the Project Files")
# ═════════════════════════════════════════════════════════════════════════════
body("You will receive the project as a ZIP archive or a folder named TerraAlert-updatedFiles. Place it anywhere on your PC — for example:")
code_block(["C:\\Projects\\TerraAlert-updatedFiles"])
body("If you received a ZIP file, right-click it and select 'Extract All', then extract to a convenient location.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("3. Opening a Terminal in the Project Folder")
# ═════════════════════════════════════════════════════════════════════════════
body("All commands below must be run inside the project folder.")
h("Windows", level=3)
numbered("Open File Explorer and navigate to the TerraAlert-updatedFiles folder.")
numbered("Click on the address bar at the top, type 'powershell', and press Enter.")
numbered("A blue PowerShell window will open, already pointing to the correct folder.")
h("macOS / Linux", level=3)
numbered("Open a Terminal.")
numbered("Type 'cd' followed by the full path to the folder and press Enter. Example:")
code_block(["cd /Users/yourname/Projects/TerraAlert-updatedFiles"])
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("4. Creating a Virtual Environment (Recommended)")
# ═════════════════════════════════════════════════════════════════════════════
body("A virtual environment keeps all project dependencies isolated so they do not interfere with other Python projects on your PC.")
h("Windows", level=3)
code_block([
    "python -m venv venv",
    "venv\\Scripts\\activate",
])
h("macOS / Linux", level=3)
code_block([
    "python3 -m venv venv",
    "source venv/bin/activate",
])
note("After activation, your terminal prompt will show (venv) at the beginning. This confirms the environment is active.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("5. Installing Dependencies")
# ═════════════════════════════════════════════════════════════════════════════
body("Install all required Python packages with a single command:")
code_block(["pip install -r requirements.txt scikit-learn joblib pandas"])
body("This will install Flask, Requests, Feedparser, Scikit-Learn, Pandas, and all other components the project needs. It may take 1–3 minutes depending on your internet connection.")
warn("Do not close the terminal while packages are installing.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("6. Training the Machine Learning Model")
# ═════════════════════════════════════════════════════════════════════════════
body("The project uses a built-in ML classifier. You must train it once before running the app. Run:")
code_block(["python ml/train_model.py"])
body("When it finishes, you will see two confirmation messages:")
bullet("Model saved to …/ml/models/disaster_classifier.pkl")
bullet("Metrics saved to …/ml/models/metrics.json")
note("You only need to run this step once. The trained model file is saved and reused every time the app starts.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("7. Running the Application")
# ═════════════════════════════════════════════════════════════════════════════
body("Start the Flask development server:")
code_block(["python run.py"])
body("You will see output similar to:")
code_block([
    " * Serving Flask app 'app'",
    " * Debug mode: on",
    " * Running on http://127.0.0.1:5000",
    "Press CTRL+C to quit",
])
warn("Do NOT close the terminal while using the app — doing so will stop the server.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("8. Opening the Application in Your Browser")
# ═════════════════════════════════════════════════════════════════════════════
body("Open any modern web browser (Chrome, Firefox, or Edge) and type the following address into the address bar:")
code_block(["http://127.0.0.1:5000"])
body("The TerraAlert home page will load. You can now use the full application.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("9. Using the Application")
# ═════════════════════════════════════════════════════════════════════════════
body("Once the app loads, you have access to the following pages via the top navigation bar:")

rows = [
    ("Home",      "/",           "Landing page with system overview and risk notification if a permanent location is set."),
    ("Dashboard", "/dashboard",  "Interactive map. Enter a city name or coordinates (lat,lon) and click Analyze to see risk data."),
    ("Guidance",  "/guidance",   "Disaster safety tips organized by type (flood, earthquake, etc.)."),
    ("Alerts",    "/alerts",     "Historical alert log showing all generated warnings."),
    ("Help",      "/help",       "Emergency contact information."),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdrs = tbl.rows[0].cells
for i, txt in enumerate(["Page", "URL", "Description"]):
    p = hdrs[i].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F4E79")
    hdrs[i]._tc.get_or_add_tcPr().append(shd)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for page, url, desc in rows:
    r = tbl.add_row().cells
    for i, txt in enumerate([page, url, desc]):
        p = r[i].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        if i == 1:
            run.font.name = "Courier New"

tbl.columns[0].width = Inches(1.1)
tbl.columns[1].width = Inches(1.4)
tbl.columns[2].width = Inches(4.0)
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("10. Stopping the Server")
# ═════════════════════════════════════════════════════════════════════════════
body("When you are finished using TerraAlert, go back to the terminal window and press:")
code_block(["CTRL + C"])
body("This will stop the Flask server cleanly. You can restart it any time with 'python run.py'.")
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("11. Troubleshooting")
# ═════════════════════════════════════════════════════════════════════════════

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = "Table Grid"
for i, txt in enumerate(["Problem", "Solution"]):
    p = tbl2.rows[0].cells[i].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "C00000")
    tbl2.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

issues = [
    ("'python' is not recognised",              "Python is not added to PATH. Reinstall Python and check 'Add Python to PATH'."),
    ("pip install fails / no internet",          "Download the project on a machine with internet first, then copy the venv folder."),
    ("Port 5000 already in use",                "Another app is using port 5000. Run: python run.py --port=5001 and visit http://127.0.0.1:5001"),
    ("'No module named flask'",                 "The virtual environment is not activated. Run the activate command from Step 4 again."),
    ("Map does not load (blank grey box)",       "Check your internet connection — map tiles are loaded from OpenStreetMap CDN."),
    ("'Model or dataset missing' error",         "You skipped Step 6. Run 'python ml/train_model.py' first, then restart the app."),
    ("Page shows 'No reports found'",           "Normal on first use. Enter a location and click Analyze to fetch real-time data."),
]

for prob, sol in issues:
    r = tbl2.add_row().cells
    r[0].paragraphs[0].add_run(prob).font.size = Pt(10)
    r[1].paragraphs[0].add_run(sol).font.size = Pt(10)

tbl2.columns[0].width = Inches(2.5)
tbl2.columns[1].width = Inches(4.0)
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
h("12. Quick Command Reference")
# ═════════════════════════════════════════════════════════════════════════════
body("All commands you need, in order, from a fresh installation on Windows:")
code_block([
    "# 1. Navigate into the project folder",
    "cd C:\\Projects\\TerraAlert-updatedFiles",
    "",
    "# 2. Create and activate virtual environment",
    "python -m venv venv",
    "venv\\Scripts\\activate",
    "",
    "# 3. Install dependencies",
    "pip install -r requirements.txt scikit-learn joblib pandas",
    "",
    "# 4. Train the ML model (first time only)",
    "python ml/train_model.py",
    "",
    "# 5. Start the application",
    "python run.py",
    "",
    "# 6. Open browser at:",
    "http://127.0.0.1:5000",
])

doc.add_page_break()
footer_p = doc.add_paragraph("TerraAlert – Early Disaster Warning System  |  Setup Guide  |  April 2026")
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.runs[0].font.size = Pt(9)
footer_p.runs[0].italic = True
footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save("TerraAlert_Setup_Guide.docx")
print("Saved: TerraAlert_Setup_Guide.docx")
