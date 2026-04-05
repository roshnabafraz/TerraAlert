"""
insert_diagrams.py
Inserts the 11 TerraAlert diagram JPGs into TerraAlert.docx as Figures
and adds Table captions above every existing table.

Strategy:
- Preserve ALL existing content / formatting untouched.
- Add a "Diagrams" section at the END of the document containing all 11 figures.
- Name each figure "Figure X – <Title>" under the image (centered, italic).
- Add "Table X – <Title>" captions ABOVE each existing table (same style as surrounding text).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os

DIAGRAMS_DIR = r"Diagrams"
DOCX_IN      = "TerraAlert.docx"
DOCX_OUT     = "TerraAlert.docx"   # overwrite in-place

DIAGRAMS = [
    ("use_case_diagram.jpg",           "Use Case Diagram – TerraAlert System Use Cases"),
    ("system_architecture_diagram.jpg","System Architecture Diagram – High-Level MVC Architecture"),
    ("activity_diagram.jpg",           "Activity Diagram – Location Risk Analysis Flow"),
    ("sequence_diagram.jpg",           "Sequence Diagram – Permanent Location Loading and Warning Intercept"),
    ("component_diagram.jpg",          "Component Diagram – Core Business Logic Modules"),
    ("state_machine_diagram.jpg",      "State Machine Diagram – Disaster Event Lifecycle"),
    ("class_diagram.jpg",              "Class Diagram – Domain Entity Model"),
    ("data_flow_diagram.jpg",          "Data Flow Diagram – Level 1 Application Data Flow"),
    ("er_diagram.jpg",                 "ER Diagram – TerraAlert SQLite Database Schema"),
    ("deployment_diagram.jpg",         "Deployment Diagram – Cloud Deployment Architecture"),
    ("ml_pipeline_diagram.jpg",        "ML Pipeline Diagram – Training and Inference Flow"),
]

TABLE_NAMES = [
    "Table 1 – Software Requirements (Functional and Non-Functional)",
    "Table 2 – API Endpoints Summary",
    "Table 3 – Database Schema: disaster_reports, alerts, guidance",
]

# ── helpers ────────────────────────────────────────────────────────────────

def add_caption(doc, text):
    """Adds a centred italic caption paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    return p


def add_table_caption_before(doc, table_index, caption_text):
    """
    Inserts a bold caption paragraph directly before the given table in the
    document body XML so that it appears ABOVE the table.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    tbl_element = doc.tables[table_index]._tbl
    body = doc.element.body

    # Create a new paragraph element
    new_para = doc.add_paragraph()  # appended at end initially
    body.remove(new_para._p)        # detach it

    run = new_para.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)

    # Re-insert before the table
    tbl_element.addprevious(new_para._p)
    return new_para


# ── main ───────────────────────────────────────────────────────────────────

doc = Document(DOCX_IN)

# 1. Add Table captions above each table
for idx, name in enumerate(TABLE_NAMES):
    if idx < len(doc.tables):
        add_table_caption_before(doc, idx, name)
        print(f"Added caption for table {idx}: {name}")

# 2. Add a page break then the Diagrams section at the end
doc.add_page_break()

heading = doc.add_heading("Diagrams", level=1)

for fig_num, (filename, title) in enumerate(DIAGRAMS, start=1):
    img_path = os.path.join(DIAGRAMS_DIR, filename)
    if not os.path.exists(img_path):
        print(f"  MISSING: {img_path} – skipping")
        continue

    # Blank line before each figure for breathing room
    doc.add_paragraph()

    # Insert image centred
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_path, width=Inches(5.5))   # fits A4 with margins

    # Caption below
    caption = add_caption(doc, f"Figure {fig_num} – {title}")
    print(f"  Inserted: Figure {fig_num} – {title}")

doc.save(DOCX_OUT)
print(f"\nSaved: {DOCX_OUT}")
