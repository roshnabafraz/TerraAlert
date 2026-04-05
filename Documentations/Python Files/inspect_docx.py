from docx import Document

doc = Document("TerraAlert.docx")

print("=== ALL NON-EMPTY PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name
    text = p.text.strip()
    if text:
        print(f"[{i}] Style='{style}' | '{text[:100]}'")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
